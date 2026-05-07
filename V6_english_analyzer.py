# -*- coding: utf-8 -*-

import re
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from openai import OpenAI
import time
import threading
from datetime import datetime
import json
from typing import Dict, List, Optional

# 配置参数
CONFIG = {
    "vocab_book_path": "vocabulary_book.json",  # 生词本保存位置
    "parsing_error_log_dir": "parsing_errors",   # 解析错误日志目录
    "api_retry_max": 3,                         # API最大重试次数
    "api_initial_timeout": 15,                  # API初始超时时间(秒)
    "api_max_timeout": 60                       # API最大超时时间(秒)
}

# 初始化DeepSeek客户端
client = OpenAI(
    base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"),
    api_key=os.getenv("DEEPSEEK_API_KEY", "")
)

class VocabularyBook:
    """生词本管理类，支持单词和词组，固定三列结构"""
    
    def __init__(self, save_path: Optional[str] = None):
        self.save_path = save_path or CONFIG["vocab_book_path"]
        self.vocabularies = self._load_vocabularies()
        
    def _load_vocabularies(self) -> List[Dict]:
        """加载已保存的生词"""
        if os.path.exists(self.save_path):
            try:
                with open(self.save_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"加载生词本失败: {e}")
                return []
        return []
    
    def save_vocabulary(self) -> None:
        """保存生词本到文件"""
        try:
            # 创建目录(如果需要)
            dir_name = os.path.dirname(self.save_path)
            if dir_name and not os.path.exists(dir_name):
                os.makedirs(dir_name)
                
            with open(self.save_path, 'w', encoding='utf-8') as f:
                json.dump(self.vocabularies, f, ensure_ascii=False, indent=2)
            print(f"生词本已保存至: {self.save_path}")
        except Exception as e:
            print(f"保存生词本失败: {e}")
    
    def add_word(self, term: str, part_of_speech: str, translation: str,
                example: str, cefr_level: str, sentence: str) -> bool:
        """添加单词或词组到生词本"""
        # 检查是否已存在（不区分大小写）
        term_lower = term.lower()
        for vocab in self.vocabularies:
            if vocab["term"].lower() == term_lower:
                print(f"{'单词' if ' ' not in term else '词组'} '{term}' 已在生词本中")
                return False
                
        # 添加新条目
        new_entry = {
            "term": term,                  # 单词或词组
            "part_of_speech": part_of_speech,  # 词性
            "translation": translation,    # 翻译
            "example": example,            # 例句
            "cefr_level": cefr_level,      # 难度级别
            "original_sentence": sentence, # 原句
            "added_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "mastered": False,
            "review_count": 0,
            "last_review": None
        }
        
        self.vocabularies.append(new_entry)
        self.save_vocabulary()
        print(f"已添加{'单词' if ' ' not in term else '词组'}: {term}")
        return True
    
    def mark_as_mastered(self, term: str) -> bool:
        """标记单词/词组为已掌握"""
        for vocab in self.vocabularies:
            if vocab["term"].lower() == term.lower():
                vocab["mastered"] = True
                vocab["last_review"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                self.save_vocabulary()
                return True
        return False
    
    def get_stats(self) -> Dict:
        """获取生词本统计信息"""
        total = len(self.vocabularies)
        mastered = sum(1 for vocab in self.vocabularies if vocab["mastered"])
        by_level = {}
        words_count = 0
        phrases_count = 0
        
        for vocab in self.vocabularies:
            # 区分单词和词组
            if ' ' in vocab["term"]:
                phrases_count += 1
            else:
                words_count += 1
                
            level = vocab["cefr_level"]
            if level not in by_level:
                by_level[level] = 0
            by_level[level] += 1
            
        return {
            "total": total,
            "words": words_count,
            "phrases": phrases_count,
            "mastered": mastered,
            "by_level": by_level,
            "recent_additions": sorted(
                self.vocabularies, 
                key=lambda x: x["added_date"], 
                reverse=True
            )[:5]
        }

def split_text_into_sentences(text):
    """
    使用正则表达式拆分文章为句子
    """
    print("拆分文章为句子中...")
    
    # 预处理文本，保护特殊情况
    protected_text = text
    
    # 扩展缩写词列表
    abbreviations = [
        'Mr.', 'Mrs.', 'Ms.', 'Dr.', 'Prof.', 'Sr.', 'Jr.', 'vs.',
        'i.e.', 'e.g.', 'etc.', 'viz.', 'cf.', 'St.', 'Ph.D.',
        'U.S.', 'U.K.', 'B.A.', 'M.A.', 'a.m.', 'p.m.', 'Inc.',
        'Corp.', 'Ltd.', 'No.', 'Vol.', 'pp.', 'et al.', 'ex.'
    ]
    
    for abbr in abbreviations:
        protected_text = protected_text.replace(abbr, abbr.replace('.', '<DOT>'))
    
    # 保护包含点号的数字(包括版本号、IP地址等)
    protected_text = re.sub(r'(\d+)\.(\d+)', r'\1<DOT>\2', protected_text)
    protected_text = re.sub(r'(\d+)\.(\d+)\.(\d+)', r'\1<DOT>\2<DOT>\3', protected_text)
    
    # 保护网址和邮箱
    protected_text = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', 
                          lambda m: m.group().replace('.', '<DOT>'), protected_text)
    protected_text = re.sub(r'https?://[^\s]+', 
                          lambda m: m.group().replace('.', '<DOT>'), protected_text)
    
    # 保护可能包含点号的名字（如 Elon.Musk）
    protected_text = re.sub(r'([A-Z][a-z]+)\.([A-Z][a-z]+)', 
                          r'\1<DOT>\2', protected_text)
    
    # 使用改进的句子分割正则表达式
    # 考虑问号、感叹号和省略号
    sentences = re.split(r'(?<=[.!?…])[\s\n]+(?=[A-Z«\"\'])', protected_text)
    
    # 还原被保护的点号
    restored_sentences = []
    for sentence in sentences:
        restored = sentence.replace('<DOT>', '.')
        # 清理多余的空白字符
        restored = ' '.join(restored.split())
        if restored.strip() and len(restored.split()) > 1:  # 只添加非空且多于一个单词的句子
            restored_sentences.append(restored)
    
    print(f"文章拆分完成，共有 {len(restored_sentences)} 个句子。")
    return restored_sentences

class TimeoutException(Exception):
    """自定义超时异常类"""
    pass

def api_call_wrapper(result_container, sentence, timeout):
    """API调用包装函数，用于在线程中执行"""
    try:
        # 强化的系统提示词，严格规定格式
        system_prompt = """你是中英文翻译专家，需严格按照以下格式分析英文句子，格式错误将导致解析失败：

【原句】
<英文原句，保持原样>

【难度分级】
- CEFR级别: <仅允许A1/A2/B1/B2/C1/C2>
- 难度理由: <简要说明>

【词汇解析】
<我现在有美国高中生的词汇量，请将可能我不认识的词汇/词组单独一行，格式必须为：>
- <术语> <词性> <翻译>。例句：<例句>
<术语可以是单词或词组，词性用缩写（如n./v./adj.），翻译和例句不可省略>

【语法分析】
<自由文本>

【中文翻译】
<流畅翻译>

【学习提示】
<自由文本>

注意：
1. 【词汇解析】中每个条目必须严格遵循"- 术语 词性 翻译。例句：例句"格式
2. 术语包含空格（词组）时不可拆分，必须作为整体出现
3. 词性部分仅允许字母、点号和斜杠（如v./n.）
4. 翻译部分以句号结尾，之后紧跟"例句："
"""

        completion = client.chat.completions.create(
            model=os.getenv("DEEPSEEK_MODEL", "deepseek-v4-flash"),
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": sentence
                }
            ],
            timeout=timeout  # 设置OpenAI客户端超时
        )
        
        result_container['result'] = completion.choices[0].message.content
        result_container['success'] = True
            
    except Exception as e:
        result_container['result'] = f"分析失败: {str(e)}"
        result_container['success'] = False

def analyze_sentence_with_timeout(sentence, timeout=20):
    """
    使用线程实现的跨平台超时控制API调用
    """
    result_container = {'result': '', 'success': False}
    
    # 创建API调用线程
    api_thread = threading.Thread(
        target=api_call_wrapper,
        args=(result_container, sentence, timeout)
    )
    
    api_thread.start()
    api_thread.join(timeout)
    
    if api_thread.is_alive():
        # 线程仍在运行，说明超时
        return "分析失败: API调用超时"
    else:
        return result_container.get('result', "分析失败: 未知错误")

def validate_api_response(result: str) -> bool:
    """校验API返回是否符合基本格式要求"""
    required_sections = ["【原句】", "【词汇解析】", "【中文翻译】"]
    for section in required_sections:
        if section not in result:
            print(f"API返回缺少必要段落：{section}")
            return False
    
    # 检查词汇解析部分是否有至少一个有效条目
    vocab_lines = [line for line in result.split('\n') if line.strip().startswith('-')]
    if len(vocab_lines) == 0:
        print("API返回的词汇解析为空")
        return False
    
    return True

def log_parsing_error(sentence: str, result: str, error: str):
    """记录解析错误日志，便于后续分析"""
    log_dir = CONFIG["parsing_error_log_dir"]
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, f"error_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
    
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(f"时间: {datetime.now()}\n")
        f.write(f"原句: {sentence}\n")
        f.write(f"错误原因: {error}\n")
        f.write(f"API返回:\n{result}\n")
    
    print(f"解析错误已记录至: {log_path}")

def analyze_sentence(sentence):
    """
    改进的句子分析函数，包含智能重试和超时控制
    """
    print(f"正在调用 API 分析句子: {sentence[:50]}...")  # 只显示前50个字符
    
    for attempt in range(CONFIG["api_retry_max"]):
        # 动态调整超时时间（指数退避）
        current_timeout = min(
            CONFIG["api_initial_timeout"] * (2** attempt), 
            CONFIG["api_max_timeout"]
        )
        
        try:
            result = analyze_sentence_with_timeout(sentence, current_timeout)
            
            # 校验返回格式
            if "分析失败" not in result and validate_api_response(result):
                return result
            else:
                print(f"API返回格式无效 (尝试 {attempt+1}/{CONFIG['api_retry_max']})")
                
        except Exception as e:
            print(f"API调用异常 (尝试 {attempt+1}/{CONFIG['api_retry_max']}): {str(e)}")
            
        # 如果不是最后一次尝试，添加延迟
        if attempt < CONFIG["api_retry_max"] - 1:
            delay = 1 + attempt  # 递增延迟时间
            print(f"等待 {delay} 秒后重试...")
            time.sleep(delay)
    
    error_msg = f"分析失败: 已达到最大重试次数 ({CONFIG['api_retry_max']}次)"
    return error_msg

def parse_api_result(result):
    """解析API返回的文本为结构化的字典"""
    sections = {
        "原句": "",
        "难度分级": "",
        "词汇解析": "",
        "语法分析": "",
        "中文翻译": "",
        "学习提示": "",
        "cefr_level": ""  # 单独提取CEFR级别
    }
    
    current_section = None
    for line in result.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # 检测章节标题
        if "【原句】" in line:
            current_section = "原句"
            continue
        elif "【难度分级】" in line:
            current_section = "难度分级"
            continue
        elif "【词汇解析】" in line:
            current_section = "词汇解析"
            continue
        elif "【语法分析】" in line:
            current_section = "语法分析"
            continue
        elif "【中文翻译】" in line:
            current_section = "中文翻译"
            continue
        elif "【学习提示】" in line:
            current_section = "学习提示"
            continue
            
        # 添加到相应章节
        if current_section and current_section in sections:
            if sections[current_section]:
                sections[current_section] += "\n" + line
            else:
                sections[current_section] = line
                
        # 提取CEFR级别
        if current_section == "难度分级" and "CEFR级别:" in line:
            cefr_match = re.search(r'CEFR级别: ([A-Z]\d)', line)
            if cefr_match:
                sections["cefr_level"] = cefr_match.group(1)
                
    return sections

def extract_vocabulary_from_analysis(analysis: Dict, sentence: str, result: str) -> List[Dict]:
    """从分析结果中提取词汇信息，支持词组识别，增强容错性"""
    vocabularies = []
    vocab_section = analysis.get("词汇解析", "")
    if not vocab_section:
        log_parsing_error(sentence, result, "词汇解析部分为空")
        return vocabularies
        
    # 按行处理，过滤空行和非词汇行
    lines = [line.strip() for line in vocab_section.split('\n') if line.strip() and line.startswith('-')]

    for line in lines:
        # 阶段1：清除前缀干扰（如"- 重要词汇1: "）
        cleaned_line = re.sub(r'^-+\s*(重要|词汇)?\d*:?\s*', '', line)
        
        # 阶段2：尝试多种格式匹配（优先级从高到低）
        # 格式1：标准格式（带例句）- 修正词性匹配，不允许空格
        match = re.match(
            r'^(.+?)\s+([a-zA-Z./&]+?)\s+(.+?)\s*。?\s*例句[:：]\s*(.+)$',
            cleaned_line
        )
        if match:
            term, pos, trans, example = match.groups()
        else:
            # 格式2：无例句但有明确词性 - 修正词性匹配，不允许空格
            match = re.match(
                r'^(.+?)\s+([a-zA-Z./&]+?)\s+(.+)$',
                cleaned_line
            )
            if match:
                term, pos, trans = match.groups()
                example = ""
            else:
                # 格式3：仅提取术语和翻译（最后的容错）
                parts = re.split(r'\s{2,}', cleaned_line, maxsplit=2)  # 用多个空格作为分隔符
                if len(parts) >= 2:
                    term, trans = parts[0], parts[1]
                    pos = "未知"
                    example = ""
                else:
                    log_parsing_error(sentence, result, f"无法解析词汇行: {line}")
                    continue  # 完全无法解析则跳过

        # 清洗提取结果
        term = term.strip()
        pos = pos.strip()
        trans = trans.strip().rstrip('。,.')  # 移除结尾标点
        
        vocabularies.append({
            "term": term,
            "part_of_speech": pos,
            "translation": trans,
            "example": example.strip()
        })

    return vocabularies

def get_processed_sentences(doc_path):
    print("检查已处理的句子...")
    if not os.path.exists(doc_path):
        return []

    doc = Document(doc_path)
    processed_sentences = []
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("英文原句:"):
            # 提取句子内容（去掉"英文原句: "前缀）
            sentence = paragraph.text.split(": ", 1)[-1]
            processed_sentences.append(sentence)
    print(f"已处理句子数: {len(processed_sentences)}")
    return processed_sentences

def append_to_word(doc_path, sentence, result, idx, vocab_book):
    if not os.path.exists(doc_path):
        doc = Document()
        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # 添加标题
        heading = doc.add_heading('英语句子分析报告', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
            
        # 添加说明段落
        intro = doc.add_paragraph()
        intro_run = intro.add_run("本文档包含英文句子分析，包括难度分级、词汇解析、语法分析和中文翻译。")
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc = Document(doc_path)
    
    # 添加分页符（每5个句子一页）
    if idx > 1 and (idx-1) % 5 == 0:
        doc.add_page_break()
    
    # 添加句子标题
    heading = doc.add_heading(f"句子 {idx}", level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 添加原句
    para_sentence = doc.add_paragraph()
    sentence_run = para_sentence.add_run("英文原句: ")
    sentence_run.bold = True
    sentence_run.font.color.rgb = RGBColor(0, 0, 139)
    
    sentence_text = para_sentence.add_run(sentence)
    sentence_text.italic = True
    
    # 解析并格式化API返回的结果
    if "分析失败" not in result:
        sections = parse_api_result(result)
        cefr_level = sections.get("cefr_level", "未知")
        
        # 提取并保存词汇到生词本
        vocabularies = extract_vocabulary_from_analysis(sections, sentence, result)
        for vocab in vocabularies:
            vocab_book.add_word(
                term=vocab["term"],
                part_of_speech=vocab["part_of_speech"],
                translation=vocab["translation"],
                example=vocab["example"],
                cefr_level=cefr_level,
                sentence=sentence
            )
        
        for section_title, section_content in sections.items():
            if section_title == "cefr_level":  # 跳过内部使用的CEFR级别字段
                continue
                
            if section_content and section_title != "原句":  # 原句已单独处理
                # 添加小节标题
                para_title = doc.add_paragraph()
                title_run = para_title.add_run(section_title + ":")
                title_run.bold = True
                if section_title == "难度分级":
                    title_run.font.color.rgb = RGBColor(153, 0, 0)  # 深红色
                elif section_title == "词汇解析":
                    title_run.font.color.rgb = RGBColor(0, 75, 0)  # 深绿色
                elif section_title == "语法分析":
                    title_run.font.color.rgb = RGBColor(102, 51, 0)  # 棕色
                elif section_title == "中文翻译":
                    title_run.font.color.rgb = RGBColor(0, 0, 153)  # 深蓝色
                elif section_title == "学习提示":
                    title_run.font.color.rgb = RGBColor(102, 0, 102)  # 紫色
                
                # 添加内容
                para_content = doc.add_paragraph(section_content)
                if section_title == "中文翻译":
                    para_content.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in para_content.runs:
                        run.font.name = '仿宋'
    else:
        # 处理分析失败的情况
        error_para = doc.add_paragraph()
        error_run = error_para.add_run("分析失败: ")
        error_run.bold = True
        error_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        error_para.add_run(result)
    
    # 添加句子分割线
    separator = doc.add_paragraph()
    separator_run = separator.add_run('=' * 50)
    separator_run.font.color.rgb = RGBColor(200, 200, 200)
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(doc_path)
    print(f"句子 {idx} 分析结果已保存到 Word 文档。")

def add_summary(doc_path, total_sentences, successful_count, failed_count, vocab_stats):
    """在文档末尾添加学习统计和生词本总结"""
    doc = Document(doc_path)
    
    # 添加总结标题
    doc.add_page_break()
    heading = doc.add_heading('学习总结', level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(102, 0, 102)  # 紫色
    
    # 添加统计信息
    stats = doc.add_paragraph()
    stats.add_run(f"分析完成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n")
    stats.add_run(f"总共分析句子: {total_sentences}个\n")
    stats.add_run(f"成功分析: {successful_count}个\n")
    stats.add_run(f"分析失败: {failed_count}个\n\n")
    
    # 计算成功率
    if total_sentences > 0:
        success_rate = (successful_count / total_sentences) * 100
        stats.add_run(f"分析成功率: {success_rate:.1f}%\n\n")
    
    # 添加生词本统计
    stats.add_run("生词本统计:\n").bold = True
    stats.add_run(f"- 总词汇数: {vocab_stats['total']}\n")
    stats.add_run(f"- 其中单词: {vocab_stats['words']}\n")
    stats.add_run(f"- 其中词组: {vocab_stats['phrases']}\n")
    stats.add_run(f"- 已掌握: {vocab_stats['mastered']}\n")
    stats.add_run("- 按难度分级:\n")
    for level, count in vocab_stats['by_level'].items():
        stats.add_run(f"  - {level}: {count}个\n")
    
    # 添加学习建议
    tips = doc.add_paragraph()
    tips.add_run("学习建议:").bold = True
    tips.add_run("\n• 每日复习已分析的句子\n• 重点关注语法分析部分\n• 尝试使用新学词汇造句\n• 每周回顾一次学习总结")
    
    # 如果有失败的句子，添加提示
    if failed_count > 0:
        failed_note = doc.add_paragraph()
        failed_note.add_run(f"\n注意: 有 {failed_count} 个句子分析失败，建议检查网络连接或API密钥。")
        failed_note.runs[-1].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.save(doc_path)

def process_text_file(txt_path, vocab_book_path=None):
    """
    输入一个txt文件路径，输出同名的Word文档，其中包含分析结果。
    """
    print(f"开始处理文件: {txt_path}")

    # 初始化生词本
    vocab_book = VocabularyBook(vocab_book_path)

    # 主处理逻辑
    try:
        doc_path = txt_path.replace('.txt', '.docx')
        progress_path = txt_path + '.progress'
        
        # 读取或创建进度文件
        processed_indices = set()
        if os.path.exists(progress_path):
            with open(progress_path, 'r', encoding='utf-8') as f:
                processed_indices = set(int(line.strip()) for line in f)
        
        print(f"读取文件内容: {txt_path}")
        with open(txt_path, "r", encoding="utf-8") as file:
            text = file.read()

        sentences = split_text_into_sentences(text)
        processed_sentences = get_processed_sentences(doc_path)
        
        successful_count = 0
        failed_count = 0

        for idx, sentence in enumerate(sentences, 1):
            # 检查句子是否已处理
            if sentence in processed_sentences or idx in processed_indices:
                print(f"跳过已处理的句子 {idx}")
                successful_count += 1
                continue

            print(f"开始处理句子 {idx}/{len(sentences)}")
            try:
                result = analyze_sentence(sentence)
                
                if "分析失败" not in result:
                    append_to_word(doc_path, sentence, result, idx, vocab_book)
                    # 更新进度文件
                    with open(progress_path, 'a', encoding='utf-8') as f:
                        f.write(f"{idx}\n")
                    successful_count += 1
                    print(f"句子 {idx} 分析成功")
                else:
                    print(f"句子 {idx} 分析失败，跳过: {result}")
                    failed_count += 1
                    
            except Exception as e:
                print(f"处理句子 {idx} 时发生异常，跳过: {str(e)}")
                failed_count += 1
                
            # 添加延迟，避免API限制
            time.sleep(0.5)  # 减少固定延迟时间
        
        # 获取生词本统计并添加学习总结
        vocab_stats = vocab_book.get_stats()
        add_summary(doc_path, len(sentences), successful_count, failed_count, vocab_stats)
        print(f"文件处理完成: {txt_path}")
        print(f"成功: {successful_count}个句子, 失败: {failed_count}个句子")
        
    except Exception as e:
        print(f"处理文件时发生严重错误: {e}")
        import traceback
        traceback.print_exc()
    
    return vocab_book

def process_folder(folder_path, vocab_book_path=None):
    """
    处理指定文件夹中的所有 .txt 文件。
    """
    print(f"开始处理文件夹: {folder_path}")
    
    # 初始化共享的生词本
    vocab_book = VocabularyBook(vocab_book_path)
    
    for file_name in os.listdir(folder_path):
        if file_name.endswith('.txt'):
            txt_path = os.path.join(folder_path, file_name)
            print(f"正在处理文件: {txt_path}")
            vocab_book = process_text_file(txt_path, vocab_book_path)
            print(f"文件处理完成: {txt_path}")
    print(f"文件夹处理完成: {folder_path}")
    
    # 打印生词本统计
    stats = vocab_book.get_stats()
    print("\n生词本统计:")
    print(f"总词汇数: {stats['total']}")
    print(f"其中单词: {stats['words']}")
    print(f"其中词组: {stats['phrases']}")
    print(f"已掌握: {stats['mastered']}")
    print("按难度分级:")
    for level, count in stats['by_level'].items():
        print(f"  {level}: {count}个")
    
    return vocab_book

# 调用代码
if __name__ == "__main__":
    # 可以自定义生词本保存位置
    custom_vocab_path = "my_vocabulary_book.json"
    
    folder_path = "txt"  # 替换为你的文件夹路径
    process_folder(folder_path, custom_vocab_path)
    print("所有任务已完成！")
